#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""make_duthao.py — LLM cục bộ sinh DỰ THẢO công văn/báo cáo trả lời 1 văn bản đến, xuất .docx
(thể thức Nghị định 30/2020), lưu `<số>_DuThao.docx` cùng thư mục văn bản, tuỳ chọn gửi Telegram.

KHÔNG mở browser. Cần: python-docx (đã có), LLM cục bộ. Dự thảo là NHÁP để người dùng hoàn thiện.
Dùng: python scripts/make_duthao.py "<số ký hiệu>" [--send]
"""
import argparse, json, re, sys
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import fetch_vanban as F
import send_next as S

WORKDIR = F.WORKDIR
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

SYSTEM = (
    "Bạn là chuyên viên văn thư Sở Giáo dục và Đào tạo, soạn DỰ THẢO văn bản trả lời theo thể thức "
    "hành chính (Nghị định 30/2020). Dựa vào văn bản đến (kèm nội dung), trả về DUY NHẤT 1 JSON:\n"
    '{"loai":"cong_van|bao_cao","trich_yeu":"...","kinh_gui":["..."],'
    '"mo_dau":"đoạn dẫn căn cứ văn bản đến","noi_dung":["mục 1","mục 2",...],'
    '"ket":"câu kết","noi_nhan":["Như trên","Lưu: VT, ..."],"can_quyet":["điểm cần người dùng quyết"]}\n'
    "Bám đúng yêu cầu của văn bản đến; nêu rõ chỗ cần số liệu bằng [CẦN SỐ LIỆU]. KHÔNG bịa số liệu. "
    "Tiếng Việt hành chính, ngắn gọn."
)


def gen(doc) -> dict:
    text = " ".join(" ".join((e.get("text") or "") for e in doc.get("extracted", [])).split())[:3000]
    user = json.dumps({"so_ky_hieu": doc.get("so_ky_hieu"), "trich_yeu": doc.get("trich_yeu"),
                       "vai_tro": doc.get("nhan_xu_ly"), "noi_gui": doc.get("noi_gui"),
                       "noi_dung_van_ban_den": text or "(chưa trích được)"}, ensure_ascii=False)
    raw = S.llm_chat([{"role": "system", "content": SYSTEM}, {"role": "user", "content": user}], max_tokens=2000)
    m = re.search(r"\{.*\}", raw, re.S)
    return json.loads(m.group(0)) if m else {"loai": "cong_van", "trich_yeu": doc.get("trich_yeu"),
                                             "kinh_gui": ["..."], "mo_dau": raw[:400], "noi_dung": [],
                                             "ket": "", "noi_nhan": ["Lưu: VT."], "can_quyet": []}


def build_docx(d: dict, so_goc: str, out_path: Path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(13)

    def p(text="", align=AL.LEFT, bold=False, size=13, italic=False):
        par = doc.add_paragraph(); par.alignment = align
        r = par.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        r.font.name = "Times New Roman"
        return par

    # Khối đầu (giản lược 1 cột — người dùng căn lại 2 cột khi hoàn thiện)
    p("SỞ GIÁO DỤC VÀ ĐÀO TẠO TỈNH TUYÊN QUANG", AL.CENTER, bold=True, size=12)
    p("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM", AL.CENTER, bold=True, size=12)
    p("Độc lập - Tự do - Hạnh phúc", AL.CENTER, bold=True, size=13)
    so = f"Số:        /{'BC-SGDĐT' if d.get('loai') == 'bao_cao' else 'SGDĐT'}"
    p(so + "          " + f"Tuyên Quang, ngày … tháng … năm {date.today().year}", AL.CENTER, italic=True, size=13)
    p()
    if d.get("loai") == "bao_cao":
        p("BÁO CÁO", AL.CENTER, bold=True, size=14)
        p(d.get("trich_yeu", ""), AL.CENTER, bold=True, size=13)
    else:
        p("V/v " + d.get("trich_yeu", ""), AL.CENTER, italic=True, size=13)
    p()
    p("Kính gửi: " + "; ".join(d.get("kinh_gui", []) or ["…"]) + ".", size=13)
    p()
    if d.get("mo_dau"):
        p(d["mo_dau"], align=AL.JUSTIFY)
    for i, mc in enumerate(d.get("noi_dung", []), 1):
        mc = re.sub(r"^\s*\d+[.)]\s*", "", mc)        # bỏ số LLM tự đánh (tránh "1. 1.")
        p(f"{i}. {mc}", align=AL.JUSTIFY)
    if d.get("ket"):
        p(); p(d["ket"], align=AL.JUSTIFY)
    p()
    # Nơi nhận + ký
    p("Nơi nhận:", bold=True, italic=True, size=12)
    for nn in (d.get("noi_nhan") or ["Lưu: VT."]):
        p("- " + nn + ";", size=12)
    p("GIÁM ĐỐC", AL.RIGHT, bold=True)
    p("(ký, đóng dấu)", AL.RIGHT, italic=True, size=12)
    p()
    if d.get("can_quyet"):
        p(">>> CẦN NGƯỜI DÙNG QUYẾT:", bold=True, size=12)
        for cq in d["can_quyet"]:
            p("   - " + cq, size=12)
    doc.save(str(out_path))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("so_ky_hieu")
    ap.add_argument("--send", action="store_true", help="Gửi file dự thảo lên Telegram")
    a = ap.parse_args()

    inbox = json.loads((WORKDIR / "inbox.json").read_text(encoding="utf-8"))
    doc = next((x for x in inbox if x.get("so_ky_hieu") == a.so_ky_hieu), None)
    if not doc:
        sys.exit(f"Không thấy văn bản {a.so_ky_hieu}")
    folder = None
    for at in doc.get("attachments", []):
        if Path(at).parent.exists():
            folder = Path(at).parent; break
    folder = folder or (F.ATTACH_DIR / F.safe_name(a.so_ky_hieu))
    folder.mkdir(parents=True, exist_ok=True)

    print(f">>> LLM soạn dự thảo cho {a.so_ky_hieu} ...")
    d = gen(doc)
    out = folder / (F.safe_name(a.so_ky_hieu) + "_DuThao.docx")
    build_docx(d, a.so_ky_hieu, out)
    print(f">>> Đã lưu {out.relative_to(WORKDIR)} (loại: {d.get('loai')}).")

    if a.send:
        import telegram_send as T
        cfg = T.load_cfg()
        T.send_document(cfg, out, caption=f"📝 Dự thảo (NHÁP) trả lời {a.so_ky_hieu} — cần duyệt/hoàn thiện.")
        print(">>> Đã gửi dự thảo lên Telegram.")


if __name__ == "__main__":
    main()
