#!/usr/bin/env python3
"""Fill a DOCX template with JSON or Excel data.

Usage:
    python fill_docx_template.py template.docx output.docx --data-file data.json
    python fill_docx_template.py template.docx output.docx --xlsx source.xlsx --sheet Sheet1 --row 2
    python fill_docx_template.py template.docx output_dir --xlsx source.xlsx --sheet Sheet1 --all-rows
"""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from openpyxl import load_workbook
from optimize_docx_layout import optimize
from create_nd30_docx import ensure_page_numbers

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_\-]+)\s*\}\}")


def load_data_from_json(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("JSON root must be an object")
    return data


def load_row_from_xlsx(path: Path, sheet_name: str, row_number: int) -> dict[str, Any]:
    wb = load_workbook(path, data_only=True)
    ws = wb[sheet_name]
    headers = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]
    values = [c.value for c in ws[row_number]]
    return {headers[i]: ("" if values[i] is None else values[i]) for i in range(min(len(headers), len(values))) if headers[i]}


def iter_rows_from_xlsx(path: Path, sheet_name: str) -> Iterable[tuple[int, dict[str, Any]]]:
    wb = load_workbook(path, data_only=True)
    ws = wb[sheet_name]
    headers = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]
    for row_idx in range(2, ws.max_row + 1):
        values = [c.value for c in ws[row_idx]]
        data = {headers[i]: ("" if values[i] is None else values[i]) for i in range(min(len(headers), len(values))) if headers[i]}
        if any(v not in ("", None) for v in data.values()):
            yield row_idx, data


def scalarize(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(str(x) for x in value)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def replace_text(text: str, data: dict[str, Any]) -> str:
    lookup = {str(k).lower(): v for k, v in data.items()}
    def repl(match: re.Match[str]) -> str:
        key = match.group(1)
        value = lookup.get(key.lower(), data.get(key, ""))
        if isinstance(value, str) and key.lower().endswith('date') and re.fullmatch(r"\d{4}-\d{2}-\d{2}", value):
            y, m, d = value.split('-')
            value = f'ngày {d} tháng {m} năm {y}'
        return scalarize(value)
    return PLACEHOLDER_RE.sub(repl, text)


def replace_in_paragraph(para, data: dict[str, Any]) -> None:
    if '{{' not in para.text or '}}' not in para.text:
        return
    for run in para.runs:
        if '{{' in run.text and '}}' in run.text:
            run.text = replace_text(run.text, data)
    if '{{' in para.text and '}}' in para.text:
        para.text = replace_text(para.text, data)


def fill_doc(doc: Document, data: dict[str, Any]) -> None:
    for para in doc.paragraphs:
        replace_in_paragraph(para, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, data)


def slugify(value: str) -> str:
    text = re.sub(r"[^A-Za-z0-9._-]+", "-", str(value).strip())
    text = text.strip("-")
    return text or "document"


def fill_once(template: Path, output: Path, data: dict[str, Any]) -> None:
    doc = Document(template)
    fill_doc(doc, data)
    ensure_page_numbers(doc)
    optimize(doc, "balanced")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("template")
    parser.add_argument("output")
    parser.add_argument("--data-file")
    parser.add_argument("--xlsx")
    parser.add_argument("--sheet")
    parser.add_argument("--row", type=int)
    parser.add_argument("--all-rows", action="store_true")
    parser.add_argument("--filename-field", default="document_number")
    args = parser.parse_args()

    template = Path(args.template)
    output = Path(args.output)
    if not template.exists():
        print(f"Template not found: {template}")
        return 1

    if args.data_file:
        data = load_data_from_json(Path(args.data_file))
        fill_once(template, output, data)
        print(f"Created: {output}")
        return 0

    if args.xlsx and args.sheet:
        xlsx = Path(args.xlsx)
        if args.all_rows:
            output.mkdir(parents=True, exist_ok=True)
            count = 0
            for row_idx, data in iter_rows_from_xlsx(xlsx, args.sheet):
                filename_seed = data.get(args.filename_field) or f"row-{row_idx}"
                out_file = output / f"{slugify(filename_seed)}.docx"
                fill_once(template, out_file, data)
                count += 1
            print(f"Created {count} documents in {output}")
            return 0
        if args.row is None:
            print("--row is required unless --all-rows is used")
            return 2
        data = load_row_from_xlsx(xlsx, args.sheet, args.row)
        fill_once(template, output, data)
        print(f"Created: {output}")
        return 0

    print("Provide either --data-file or --xlsx with --sheet")
    return 3


if __name__ == "__main__":
    raise SystemExit(main())
