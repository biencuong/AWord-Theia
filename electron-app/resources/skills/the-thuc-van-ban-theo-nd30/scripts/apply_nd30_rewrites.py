#!/usr/bin/env python3
from __future__ import annotations
import argparse, html, json, re, shutil
from pathlib import Path
from typing import Any
from docx import Document

def load_json(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding='utf-8'))
    if not isinstance(data, dict):
        raise ValueError('JSON root must be an object')
    return data

def select_items(report: dict[str, Any], select: str | None, all_safe: bool) -> list[dict[str, Any]]:
    items = report.get('rewrite_suggestions', [])
    for i, item in enumerate(items, 1): item.setdefault('id', i)
    if all_safe:
        return [x for x in items if x.get('change_type') == 'safe_rewrite']
    if not select: return []
    wanted = {int(x.strip()) for x in select.split(',') if x.strip().isdigit()}
    return [x for x in items if int(x.get('id', 0)) in wanted]

def iter_paragraphs(doc: Document):
    for p in doc.paragraphs: yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: yield p

def try_replace_within_run(paragraph, original: str, rewritten: str) -> bool:
    for run in paragraph.runs:
        if original in run.text:
            run.text = run.text.replace(original, rewritten)
            return True
    return False

def rebuild_paragraph_with_first_style(paragraph, new_text: str) -> bool:
    if not paragraph.runs:
        paragraph.text = new_text
        return True
    first = paragraph.runs[0]
    font_name, font_size, bold, italic, underline = first.font.name, first.font.size, first.bold, first.italic, first.underline
    for run in paragraph.runs: run.text = ''
    paragraph.runs[0].text = new_text
    paragraph.runs[0].font.name = font_name
    paragraph.runs[0].font.size = font_size
    paragraph.runs[0].bold = bold
    paragraph.runs[0].italic = italic
    paragraph.runs[0].underline = underline
    return True

def apply_rewrite(doc: Document, original: str, rewritten: str) -> int:
    count = 0
    original_norm = original.strip()
    if not original_norm: return 0
    pattern = re.compile(re.escape(original_norm), re.IGNORECASE)
    for p in iter_paragraphs(doc):
        full = p.text or ''
        if not full or not pattern.search(full): continue
        if try_replace_within_run(p, original_norm, rewritten):
            count += 1
        else:
            rebuilt = pattern.sub(rewritten, full)
            if rebuilt != full:
                rebuild_paragraph_with_first_style(p, rebuilt)
                count += 1
    return count

def write_md(summary: dict[str, Any], path: Path) -> None:
    lines = ['# Applied ND30 safe rewrites','',f"- applied_count: {summary['applied_count']}",'','## Applied items']
    for item in summary.get('applied_items', []):
        lines.append(f"- {item['id']}. {item['original']} -> {item['rewritten']} (matches={item['applied_matches']})")
    path.write_text('\n'.join(lines), encoding='utf-8')

def write_html_preview(summary: dict[str, Any], path: Path) -> None:
    rows = []
    for item in summary.get('applied_items', []):
        rows.append("<tr><td>{id}</td><td><pre>{orig}</pre></td><td><pre>{rew}</pre></td><td>{m}</td></tr>".format(id=item['id'], orig=html.escape(item['original']), rew=html.escape(item['rewritten']), m=item['applied_matches']))
    html_doc = """<!doctype html><html><head><meta charset='utf-8'><title>ND30 Safe Rewrite Preview</title>
<style>body{{font-family:Arial,sans-serif;margin:24px}}table{{border-collapse:collapse;width:100%}}td,th{{border:1px solid #ccc;padding:8px;vertical-align:top}}pre{{white-space:pre-wrap;margin:0}}</style>
</head><body><h1>ND30 Safe Rewrite Preview</h1><p>Applied count: {count}</p><table><thead><tr><th>ID</th><th>Before</th><th>After</th><th>Matches</th></tr></thead><tbody>{rows}</tbody></table></body></html>""".format(count=summary['applied_count'], rows=''.join(rows))
    path.write_text(html_doc, encoding='utf-8')

def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument('input_docx')
    parser.add_argument('review_report_json')
    parser.add_argument('output_docx')
    parser.add_argument('--select')
    parser.add_argument('--all-safe', action='store_true')
    parser.add_argument('--json-summary')
    parser.add_argument('--md-summary')
    parser.add_argument('--html-preview')
    parser.add_argument('--bundle-dir')
    args = parser.parse_args()
    src, report_path, dst = Path(args.input_docx), Path(args.review_report_json), Path(args.output_docx)
    if not src.exists() or not report_path.exists():
        print('Input DOCX or review report not found')
        return 2
    report = load_json(report_path)
    selected = select_items(report, args.select, args.all_safe)
    doc = Document(src)
    applied_items = []
    total = 0
    for item in selected:
        original, rewritten = str(item.get('original','')).strip(), str(item.get('rewritten','')).strip()
        if not original or not rewritten or original == rewritten: continue
        matches = apply_rewrite(doc, original, rewritten)
        total += matches
        applied_items.append({'id': item.get('id'), 'original': original, 'rewritten': rewritten, 'change_type': item.get('change_type','safe_rewrite'), 'reason': item.get('reason',''), 'applied_matches': matches})
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)
    summary = {'input_docx': str(src), 'output_docx': str(dst), 'selected_count': len(selected), 'applied_count': total, 'applied_items': applied_items}
    json_path = Path(args.json_summary) if args.json_summary else None
    md_path = Path(args.md_summary) if args.md_summary else None
    html_path = Path(args.html_preview) if args.html_preview else None
    if json_path: json_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding='utf-8')
    if md_path: write_md(summary, md_path)
    if html_path: write_html_preview(summary, html_path)
    if args.bundle_dir:
        bundle = Path(args.bundle_dir)
        bundle.mkdir(parents=True, exist_ok=True)
        (bundle/'bundle_manifest.json').write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding='utf-8')
        shutil.copy2(dst, bundle/dst.name)
        if json_path and json_path.exists(): shutil.copy2(json_path, bundle/json_path.name)
        if md_path and md_path.exists(): shutil.copy2(md_path, bundle/md_path.name)
        if html_path and html_path.exists(): shutil.copy2(html_path, bundle/html_path.name)
    print(json.dumps(summary, ensure_ascii=False))
    return 0
if __name__ == '__main__':
    raise SystemExit(main())
