#!/usr/bin/env python3
"""Replicate strategy that preserves ND30 fixed blocks (quoc hieu-tieu ngu,
ten co quan, noi nhan, chu ky) byte-for-byte from a source mau, only
rebuilding the paragraphs/tables strictly BETWEEN the header table and the
"Noi nhan" table.

Day la 1 chien luoc THAY THE cho generate_replicate_shell() (shell-rebuild
mac dinh cua generate_from_profile_and_content.py) - shell-rebuild rebuild
toan bo khoi co dinh bang cac ham cung (width cot 7.2/1.2/8.0cm...) nen
KHONG giu dung ty le/dinh dang mau nguon. Dung script nay khi mau nguon can
giu nguyen TUYET DOI. Xem
references/nd30-replicate-fixed-block-preservation.md.

PHAM VI: chi thay noi dung GIUA bang header va bang "Noi nhan"/chu ky. Neu
mau nguon co noi dung phu luc/section khac SAU bang Noi nhan (vi du phu luc
kho ngang cu), noi dung do duoc GIU NGUYEN KHONG DONG - caller tu xu ly
rieng (vi du dung nd30_shape_and_measure.add_landscape_section) neu can
thay the phu luc.

Usage:
    python compose_preserve_fixed_blocks.py source-template.docx content.json output.docx
"""
from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

sys.path.insert(0, str(Path(__file__).resolve().parent))
from nd30_docx_tools import reset_section_page_setup  # noqa: E402
from nd30_shape_and_measure import fix_orphan_if_present  # noqa: E402


ALIGN_MAP = {
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def locate_fixed_block_boundaries(doc: Document) -> dict[str, Any]:
    """Tim ranh gioi giua khoi header co dinh (bang dau tien trong body -
    thuong la bang 3-vung/2-vung chua quoc hieu-tieu ngu + ten co quan) va
    khoi cuoi co dinh (bang co cell dau chua 'Noi nhan' - thuong cung chua
    chu ky). Heuristic don gian nhung dung duoc voi da so mau bao cao/cong
    van chuan NĐ30. Neu khong chac chan, tra ve ambiguous=True va KHONG
    doan bua - caller phai truyen fixed_block_markers thu cong qua
    compose_preserving_fixed_blocks()."""
    body = doc.element.body
    children = list(body)
    tables = [(idx, Table(child, doc)) for idx, child in enumerate(children) if isinstance(child, CT_Tbl)]

    if not tables:
        return {'ambiguous': True, 'reason': 'Khong tim thay bang nao trong body - khong the xac dinh khoi co dinh tu dong.'}

    head_idx, _head_tbl = tables[0]

    tail_idx = None
    for idx, tbl in tables:
        if idx == head_idx:
            continue
        try:
            first_cell_text = tbl.rows[0].cells[0].text
        except Exception:
            continue
        norm = first_cell_text.strip().lower()
        if 'nơi nhận' in norm or 'noi nhan' in norm:
            tail_idx = idx
            break

    if tail_idx is None:
        if len(tables) >= 2:
            tail_idx = tables[-1][0]
        else:
            return {
                'ambiguous': True,
                'reason': 'Chi co 1 bang duy nhat trong body va khong co bang nao chua '
                          '"Noi nhan" - khong du de xac dinh ranh gioi than bai an toan.',
                'head_table_index': head_idx,
            }

    if tail_idx <= head_idx:
        return {
            'ambiguous': True,
            'reason': 'Bang duoc doan la khoi Noi nhan/chu ky khong nam sau bang header - '
                      'cau truc khong khop gia dinh mac dinh.',
            'head_table_index': head_idx,
            'tail_table_index': tail_idx,
        }

    body_start_index = head_idx + 1
    body_end_index = tail_idx - 1
    if body_start_index > body_end_index:
        return {
            'ambiguous': True,
            'reason': 'Khong co khoang trong giua bang header va bang Noi nhan de coi la '
                      'than bai can thay noi dung.',
            'head_table_index': head_idx,
            'tail_table_index': tail_idx,
        }

    return {
        'ambiguous': False,
        'head_table_index': head_idx,
        'tail_table_index': tail_idx,
        'body_start_index': body_start_index,
        'body_end_index': body_end_index,
    }


def _find_body_paragraph_index_by_text(doc: Document, substring: str) -> int | None:
    body = doc.element.body
    for idx, child in enumerate(body):
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            if substring in p.text:
                return idx
    return None


def _resolve_boundaries(doc: Document, fixed_block_markers: dict[str, str] | None) -> dict[str, Any]:
    auto = locate_fixed_block_boundaries(doc)
    if not fixed_block_markers:
        return auto

    result = dict(auto)
    body_start_marker = fixed_block_markers.get('body_start_after')
    body_end_marker = fixed_block_markers.get('body_end_before')
    if body_start_marker:
        idx = _find_body_paragraph_index_by_text(doc, body_start_marker)
        if idx is not None:
            result['body_start_index'] = idx + 1
    if body_end_marker:
        idx = _find_body_paragraph_index_by_text(doc, body_end_marker)
        if idx is not None:
            result['body_end_index'] = idx - 1
    if 'body_start_index' in result and 'body_end_index' in result:
        result['ambiguous'] = result['body_start_index'] > result['body_end_index']
    return result


def _iter_all_paragraphs_with_body_index(doc: Document):
    """Duyet moi paragraph trong doc (ke ca trong bang), gan them body_index
    = vi tri cua PHAN TU GOC (paragraph hoac bang chua no) trong body, de
    biet no co nam trong khoang bi rebuild hay khong."""
    body = doc.element.body
    for top_idx, child in enumerate(body):
        if isinstance(child, CT_P):
            yield top_idx, Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield top_idx, p


def _apply_whole_paragraph_replacements(doc: Document, boundaries: dict[str, Any],
                                         replacements: dict[str, str]) -> int:
    """Thay TOAN BO text cua 1 doan van (KHONG PHAI substring) khi
    paragraph.text.strip() KHOP CHINH XAC 1 key trong replacements - giu
    nguyen dinh dang (rPr/pPr) cua run dau tien, chi xoa text cac run con
    lai. Day la ky thuat AN TOAN da dung thanh cong trong thuc te cho cac
    dong nguyen-khoi nhu 'Tuyen Quang, ngay... thang... nam...' hoac
    'So: .../...-SGDDT'. KHONG ho tro thay the mot phan cau (substring) -
    chi ap dung khi CA DOAN VAN khop dung 1 key, tranh pha vo dinh dang cua
    cac run khac trong cung doan (bai hoc: ghi de cell.text="" da tung pha
    dinh dang 'Noi nhan:' trong thuc te)."""
    if not replacements:
        return 0
    lo = boundaries.get('body_start_index', 1 << 30)
    hi = boundaries.get('body_end_index', -1)
    count = 0
    for idx, p in _iter_all_paragraphs_with_body_index(doc):
        if lo <= idx <= hi:
            continue  # nam trong vung se bi rebuild, khong can thay o day
        key = p.text.strip()
        if key in replacements:
            runs = p.runs
            if not runs:
                continue
            runs[0].text = replacements[key]
            for r in runs[1:]:
                r.text = ''
            count += 1
    return count


def _set_font(run, *, bold=False, italic=False, size_pt=14, font_name='Times New Roman'):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _insert_body_paragraphs(doc: Document, after_element, items: list[dict[str, Any]], *,
                             usable_width_cm: float | None) -> None:
    """Chen danh sach doan van MOI ngay SAU after_element (1 phan tu XML co
    san trong body), giu dung thu tu. Ap quy tac spacing chuan (6pt ca 2
    chieu, line-spacing single - xem references/nd30-spacing-and-line-rules.md)
    va tu dong sua tu mo coi cho doan justify neu usable_width_cm duoc cung
    cap (xem references/nd30-orphan-word-control.md)."""
    anchor = after_element
    for item in items:
        p = doc.add_paragraph()
        anchor.addnext(p._p)
        anchor = p._p
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.0
        align = item.get('align', 'justify')
        p.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
        indent_cm = item.get('first_line_indent_cm')
        if indent_cm is not None:
            pf.first_line_indent = Cm(indent_cm)
        r = p.add_run(item.get('text', ''))
        _set_font(r, bold=bool(item.get('bold')), italic=bool(item.get('italic')),
                  size_pt=item.get('size_pt', 14))
        if usable_width_cm and align == 'justify':
            fix_orphan_if_present(p, column_width_cm=usable_width_cm,
                                   first_line_indent_cm=indent_cm or 0.0)


def compose_preserving_fixed_blocks(source_docx: Path, content: dict[str, Any], output: Path, *,
                                     fixed_block_markers: dict[str, str] | None = None,
                                     usable_width_cm: float | None = None) -> dict[str, Any]:
    """Chi xoa/rebuild cac paragraph/table NAM GIUA khoi header va khoi Noi
    nhan/chu ky; giu nguyen 100% moi thu truoc/sau ranh gioi do (khong doi
    width cot, table layout, hay cau truc paragraph/run cua cac khoi co
    dinh). Neu khong xac dinh duoc ranh gioi an toan, KHONG generate - tra
    ve fallback_used=True kem ly do ro rang thay vi doan bua.

    content = {
      'body': [{'text','bold','italic','align' in left/center/right/justify,
                'first_line_indent_cm','size_pt'}, ...],
      'preserved_text_replacements': {'dong van ban goc (khop chinh xac)': 'noi dung moi'}
    }
    """
    doc = Document(source_docx)

    # Chup lai kho giay/huong/le THAT SU cua section dau TRUOC khi dong gi,
    # phong truong hop vung se xoa (hiem khi, nhung co the) chua 1
    # section-break nhung - tranh loi ke thua nham sectPr da gap thuc te.
    sec0 = doc.sections[0]
    page_snapshot = {
        'orientation': 'landscape' if sec0.orientation == WD_ORIENT.LANDSCAPE else 'portrait',
        'page_width_mm': sec0.page_width.mm,
        'page_height_mm': sec0.page_height.mm,
        'top_mm': sec0.top_margin.mm,
        'bottom_mm': sec0.bottom_margin.mm,
        'left_mm': sec0.left_margin.mm,
        'right_mm': sec0.right_margin.mm,
    }

    boundaries = _resolve_boundaries(doc, fixed_block_markers)
    if boundaries.get('ambiguous'):
        return {'fallback_used': True, 'reason': boundaries.get('reason', 'unknown'), 'boundaries': boundaries}

    body = doc.element.body
    children = list(body)
    start_idx = boundaries['body_start_index']
    end_idx = boundaries['body_end_index']
    anchor = children[start_idx - 1]
    to_remove = children[start_idx:end_idx + 1]

    replacement_count = _apply_whole_paragraph_replacements(
        doc, boundaries, content.get('preserved_text_replacements', {})
    )

    for el in to_remove:
        body.remove(el)

    _insert_body_paragraphs(doc, anchor, content.get('body', []), usable_width_cm=usable_width_cm)

    reset_section_page_setup(doc.sections[0], **page_snapshot)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return {
        'fallback_used': False,
        'boundaries': boundaries,
        'body_paragraph_count': len(content.get('body', [])),
        'preserved_text_replacement_count': replacement_count,
    }


def main(argv: list[str]) -> int:
    if len(argv) < 4:
        print('Usage: python compose_preserve_fixed_blocks.py source-template.docx content.json output.docx')
        return 1
    source = Path(argv[1])
    content_path = Path(argv[2])
    output = Path(argv[3])
    if not source.exists():
        print(f'Source template not found: {source}')
        return 2
    if not content_path.exists():
        print(f'Content JSON not found: {content_path}')
        return 2

    content = json.loads(content_path.read_text(encoding='utf-8'))
    if not isinstance(content, dict):
        print('Content JSON root must be an object')
        return 3

    markers = content.get('fixed_block_markers')
    usable_width_cm = content.get('usable_width_cm')
    result = compose_preserving_fixed_blocks(
        source, content, output,
        fixed_block_markers=markers, usable_width_cm=usable_width_cm,
    )
    if result.get('fallback_used'):
        print('KHONG the xac dinh ranh gioi khoi co dinh an toan - khong generate.')
        print(f"Ly do: {result.get('reason')}")
        print('Truyen "fixed_block_markers": {"body_start_after": "...", "body_end_before": "..."} '
              'trong content JSON de chi dinh thu cong.')
        return 4
    print(f'Created {output} tu {source}, giu nguyen khoi co dinh.')
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == '__main__':
    raise SystemExit(main(sys.argv))
