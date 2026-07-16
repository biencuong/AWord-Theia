#!/usr/bin/env python3
from __future__ import annotations

import json
import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any, Iterator
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_\-]+)\s*\}\}")
DATE_LINE_RE = re.compile(r"[A-ZÀ-Ỵa-zà-ỵ .-]+,\s*ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}")
NUMBER_LINE_RE = re.compile(r"^Số:\s*[^\n]+", re.IGNORECASE | re.MULTILINE)

NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
}

EXPECTED_TITLE_BY_TYPE = {
    'cong-dien': 'CÔNG ĐIỆN',
    'giay-moi': 'GIẤY MỜI',
    'giay-gioi-thieu': 'GIẤY GIỚI THIỆU',
    'bien-ban': 'BIÊN BẢN',
    'giay-nghi-phep': 'GIẤY NGHỈ PHÉP',
    'nghi-quyet-ca-biet': 'NGHỊ QUYẾT',
    'quyet-dinh-truc-tiep': 'QUYẾT ĐỊNH',
    'quyet-dinh-gian-tiep': 'QUYẾT ĐỊNH',
}

NAMED_DOC_TITLES = [
    'THÔNG BÁO', 'KẾ HOẠCH', 'BÁO CÁO', 'TỜ TRÌNH', 'HƯỚNG DẪN',
    'PHƯƠNG ÁN', 'ĐỀ ÁN', 'DỰ ÁN', 'QUY CHẾ', 'QUY ĐỊNH', 'HỢP ĐỒNG',
    'BẢN GHI NHỚ', 'BẢN THỎA THUẬN', 'GIẤY ỦY QUYỀN', 'PHIẾU GỬI',
    'PHIẾU CHUYỂN', 'PHIẾU BÁO', 'THƯ CÔNG'
]

TOP_TITLE_EXCLUSIONS = {
    'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM',
    'ĐỘC LẬP - TỰ DO - HẠNH PHÚC',
    '________________', '________________________________________', '____________',
}


def title_candidates(entries: list[dict[str, Any]], limit: int = 24) -> list[str]:
    candidates: list[str] = []
    for entry in (entries or [])[:limit]:
        text = normalize_space(entry.get('text', ''))
        if not text:
            continue
        upper = text.upper()
        if upper in TOP_TITLE_EXCLUSIONS:
            continue
        if upper.startswith('SỐ:') or upper.startswith('V/V'):
            continue
        if DATE_LINE_RE.search(text):
            continue
        if entry.get('alignment') != WD_ALIGN_PARAGRAPH.CENTER:
            continue
        if not entry.get('all_caps') and entry.get('bold_ratio', 0) < 0.6:
            continue
        if len(text) > 90:
            continue
        candidates.append(upper)
    return candidates


def guess_named_document_title(candidates: list[str]) -> str | None:
    for title in NAMED_DOC_TITLES:
        if title in candidates:
            return title
    return None


OOXML_TEXT_PARTS = {
    'word/document.xml',
    'word/footnotes.xml',
    'word/endnotes.xml',
    'word/comments.xml',
}


def _safe(v: Any) -> str:
    return '' if v is None else str(v)


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", ' ', _safe(text)).strip()


def iter_block_items(parent: _Document | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f'Unsupported parent type: {type(parent)}')
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def iter_paragraphs_in_order(doc: Document) -> Iterator[tuple[int, Paragraph, str]]:
    seq = 0

    def walk(parent: _Document | _Cell, origin: str) -> Iterator[tuple[int, Paragraph, str]]:
        nonlocal seq
        for block in iter_block_items(parent):
            if isinstance(block, Paragraph):
                yield seq, block, origin
                seq += 1
            else:
                for row_idx, row in enumerate(block.rows):
                    for col_idx, cell in enumerate(row.cells):
                        cell_origin = f'{origin}/table[r{row_idx}c{col_idx}]'
                        yield from walk(cell, cell_origin)

    yield from walk(doc, 'body')


def _paragraph_xml_details(p: Paragraph) -> dict[str, Any]:
    pPr = p._p.pPr
    details: dict[str, Any] = {}
    if pPr is None:
        return details
    ind = pPr.find('w:ind', NS)
    if ind is not None:
        details['indent'] = {
            'left': ind.get(f"{{{NS['w']}}}left"),
            'right': ind.get(f"{{{NS['w']}}}right"),
            'firstLine': ind.get(f"{{{NS['w']}}}firstLine"),
            'hanging': ind.get(f"{{{NS['w']}}}hanging"),
        }
    spacing = pPr.find('w:spacing', NS)
    if spacing is not None:
        details['spacing'] = {
            'before': spacing.get(f"{{{NS['w']}}}before"),
            'after': spacing.get(f"{{{NS['w']}}}after"),
            'line': spacing.get(f"{{{NS['w']}}}line"),
            'lineRule': spacing.get(f"{{{NS['w']}}}lineRule"),
        }
    tabs = pPr.find('w:tabs', NS)
    if tabs is not None:
        details['tabs'] = [
            {
                'val': t.get(f"{{{NS['w']}}}val"),
                'pos': t.get(f"{{{NS['w']}}}pos"),
                'leader': t.get(f"{{{NS['w']}}}leader"),
            }
            for t in tabs.findall('w:tab', NS)
        ]
    jc = pPr.find('w:jc', NS)
    if jc is not None:
        details['jc'] = jc.get(f"{{{NS['w']}}}val")
    numPr = pPr.find('w:numPr', NS)
    if numPr is not None:
        ilvl = numPr.find('w:ilvl', NS)
        numId = numPr.find('w:numId', NS)
        details['numbering'] = {
            'ilvl': ilvl.get(f"{{{NS['w']}}}val") if ilvl is not None else None,
            'numId': numId.get(f"{{{NS['w']}}}val") if numId is not None else None,
        }
    keep_next = pPr.find('w:keepNext', NS)
    if keep_next is not None:
        details['keep_next'] = True
    page_break_before = pPr.find('w:pageBreakBefore', NS)
    if page_break_before is not None:
        details['page_break_before'] = True
    return details


def paragraph_entries(doc: Document) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for idx, p, origin in iter_paragraphs_in_order(doc):
        text = normalize_space(p.text)
        if not text:
            continue
        runs = [r for r in p.runs if normalize_space(r.text)]
        font_names: list[str] = []
        font_sizes: list[float] = []
        bold_runs = 0
        italic_runs = 0
        for r in runs:
            if r.font.name:
                font_names.append(r.font.name)
            if r.font.size:
                try:
                    font_sizes.append(round(r.font.size.pt, 1))
                except Exception:
                    pass
            if r.bold:
                bold_runs += 1
            if r.italic:
                italic_runs += 1
        xml_details = _paragraph_xml_details(p)
        entries.append({
            'index': idx,
            'origin': origin,
            'text': text,
            'alignment': p.alignment,
            'style_name': getattr(p.style, 'name', '') if p.style is not None else '',
            'font_names': sorted(set(font_names)),
            'font_sizes': sorted(set(font_sizes)),
            'all_caps': text.upper() == text and any(ch.isalpha() for ch in text),
            'contains_placeholder': bool(PLACEHOLDER_RE.search(text)),
            'bold_ratio': (bold_runs / len(runs)) if runs else 0.0,
            'italic_ratio': (italic_runs / len(runs)) if runs else 0.0,
            'xml': xml_details,
        })
    return entries


def header_footer_text(doc: Document) -> dict[str, list[str]]:
    headers: list[str] = []
    footers: list[str] = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            t = normalize_space(p.text)
            if t:
                headers.append(t)
        for p in section.footer.paragraphs:
            t = normalize_space(p.text)
            if t:
                footers.append(t)
    return {'headers': headers[:20], 'footers': footers[:20]}


def section_metrics(doc: Document) -> dict[str, Any]:
    s = doc.sections[0]
    return {
        'page_width_mm': round(s.page_width.mm, 1),
        'page_height_mm': round(s.page_height.mm, 1),
        'top_margin_mm': round(s.top_margin.mm, 1),
        'bottom_margin_mm': round(s.bottom_margin.mm, 1),
        'left_margin_mm': round(s.left_margin.mm, 1),
        'right_margin_mm': round(s.right_margin.mm, 1),
        'header_distance_mm': round(s.header_distance.mm, 1),
        'footer_distance_mm': round(s.footer_distance.mm, 1),
        'gutter_mm': round(s.gutter.mm, 1),
    }


def detect_document_type(text: str, entries: list[dict[str, Any]] | None = None) -> str:
    t = text.upper()
    top_entries = (entries or [])[:30]
    top = ' '.join(e['text'].upper() for e in top_entries)
    combined = f'{top} {t}'
    titles = title_candidates(entries or [])
    title_blob = ' | '.join(titles)

    if any(x in title_blob for x in ['SAO Y', 'SAO LỤC', 'TRÍCH SAO']):
        return 'ban-sao-giay'
    if any(x in title_blob for x in ['PHỤ LỤC']) or 'PHỤ LỤC' in combined.split('KÈM THEO')[0]:
        return 'phu-luc'
    if 'CÔNG ĐIỆN' in title_blob or 'CÔNG ĐIỆN' in combined:
        return 'cong-dien'
    if 'GIẤY MỜI' in title_blob or 'GIẤY MỜI' in combined:
        return 'giay-moi'
    if 'GIẤY GIỚI THIỆU' in title_blob or 'GIẤY GIỚI THIỆU' in combined:
        return 'giay-gioi-thieu'
    if 'BIÊN BẢN' in title_blob or 'BIÊN BẢN' in combined:
        return 'bien-ban'
    if 'GIẤY NGHỈ PHÉP' in title_blob or 'GIẤY NGHỈ PHÉP' in combined:
        return 'giay-nghi-phep'
    if 'NGHỊ QUYẾT' in title_blob or 'NGHỊ QUYẾT' in combined:
        return 'nghi-quyet-ca-biet'
    if 'QUYẾT ĐỊNH' in title_blob or 'QUYẾT ĐỊNH' in combined:
        if any(x in combined for x in ['BAN HÀNH', 'PHÊ DUYỆT']):
            return 'quyet-dinh-gian-tiep'
        return 'quyet-dinh-truc-tiep'

    named_title = guess_named_document_title(titles)
    if named_title:
        return 'van-ban-co-ten-loai'

    if 'KÍNH GỬI' in combined or 'KÍNH GỬI:' in combined:
        return 'cong-van'
    if NUMBER_LINE_RE.search(combined) and any('V/V' in e.get('text','').upper() for e in top_entries[:8]):
        return 'cong-van'

    for candidate in NAMED_DOC_TITLES:
        if candidate in combined:
            return 'van-ban-co-ten-loai'
    if any(x in combined for x in ['SAO Y', 'SAO LỤC', 'TRÍCH SAO']):
        return 'ban-sao-giay'
    return 'unknown'


def _parse_styles_xml(root: ET.Element) -> dict[str, Any]:
    styles: dict[str, Any] = {}
    for style in root.findall('.//w:style', NS):
        style_id = style.get(f"{{{NS['w']}}}styleId")
        style_type = style.get(f"{{{NS['w']}}}type")
        if not style_id:
            continue
        name_elem = style.find('w:name', NS)
        style_name = name_elem.get(f"{{{NS['w']}}}val") if name_elem is not None else None
        based_on = style.find('w:basedOn', NS)
        next_style = style.find('w:next', NS)
        linked = style.find('w:link', NS)
        ui_priority = style.find('w:uiPriority', NS)
        qformat = style.find('w:qFormat', NS) is not None
        fonts: dict[str, Any] = {}
        ppr_summary: dict[str, Any] = {}
        rpr_summary: dict[str, Any] = {}
        rpr = style.find('w:rPr', NS)
        if rpr is not None:
            font_elem = rpr.find('w:rFonts', NS)
            if font_elem is not None:
                fonts = {
                    'ascii': font_elem.get(f"{{{NS['w']}}}ascii"),
                    'hAnsi': font_elem.get(f"{{{NS['w']}}}hAnsi"),
                    'eastAsia': font_elem.get(f"{{{NS['w']}}}eastAsia"),
                    'cs': font_elem.get(f"{{{NS['w']}}}cs"),
                }
            sz_elem = rpr.find('w:sz', NS)
            if sz_elem is not None:
                fonts['size'] = sz_elem.get(f"{{{NS['w']}}}val")
            color = rpr.find('w:color', NS)
            if color is not None:
                rpr_summary['color'] = color.get(f"{{{NS['w']}}}val")
            if rpr.find('w:b', NS) is not None:
                rpr_summary['bold'] = True
            if rpr.find('w:i', NS) is not None:
                rpr_summary['italic'] = True
            if rpr.find('w:caps', NS) is not None:
                rpr_summary['caps'] = True
        ppr = style.find('w:pPr', NS)
        if ppr is not None:
            jc_elem = ppr.find('w:jc', NS)
            if jc_elem is not None:
                ppr_summary['alignment'] = jc_elem.get(f"{{{NS['w']}}}val")
            spacing_elem = ppr.find('w:spacing', NS)
            if spacing_elem is not None:
                ppr_summary['spacing'] = {
                    'before': spacing_elem.get(f"{{{NS['w']}}}before"),
                    'after': spacing_elem.get(f"{{{NS['w']}}}after"),
                    'line': spacing_elem.get(f"{{{NS['w']}}}line"),
                    'lineRule': spacing_elem.get(f"{{{NS['w']}}}lineRule"),
                }
            ind = ppr.find('w:ind', NS)
            if ind is not None:
                ppr_summary['indent'] = {
                    'left': ind.get(f"{{{NS['w']}}}left"),
                    'right': ind.get(f"{{{NS['w']}}}right"),
                    'firstLine': ind.get(f"{{{NS['w']}}}firstLine"),
                    'hanging': ind.get(f"{{{NS['w']}}}hanging"),
                }
            outline = ppr.find('w:outlineLvl', NS)
            if outline is not None:
                ppr_summary['outlineLvl'] = outline.get(f"{{{NS['w']}}}val")
        styles[style_id] = {
            'id': style_id,
            'name': style_name,
            'type': style_type,
            'basedOn': based_on.get(f"{{{NS['w']}}}val") if based_on is not None else None,
            'next': next_style.get(f"{{{NS['w']}}}val") if next_style is not None else None,
            'linked': linked.get(f"{{{NS['w']}}}val") if linked is not None else None,
            'uiPriority': ui_priority.get(f"{{{NS['w']}}}val") if ui_priority is not None else None,
            'qFormat': qformat,
            'fonts': fonts,
            'paragraph': ppr_summary,
            'run': rpr_summary,
        }
    return styles


def _parse_numbering_xml(root: ET.Element) -> dict[str, Any]:
    abstract_defs: dict[str, Any] = {}
    for abstract in root.findall('.//w:abstractNum', NS):
        abstract_id = abstract.get(f"{{{NS['w']}}}abstractNumId")
        if not abstract_id:
            continue
        levels = []
        for lvl in abstract.findall('w:lvl', NS):
            ilvl = lvl.get(f"{{{NS['w']}}}ilvl")
            numFmt = lvl.find('w:numFmt', NS)
            lvlText = lvl.find('w:lvlText', NS)
            start = lvl.find('w:start', NS)
            pPr = lvl.find('w:pPr', NS)
            entry = {
                'ilvl': ilvl,
                'numFmt': numFmt.get(f"{{{NS['w']}}}val") if numFmt is not None else None,
                'lvlText': lvlText.get(f"{{{NS['w']}}}val") if lvlText is not None else None,
                'start': start.get(f"{{{NS['w']}}}val") if start is not None else None,
            }
            if pPr is not None:
                ind = pPr.find('w:ind', NS)
                if ind is not None:
                    entry['indent'] = {
                        'left': ind.get(f"{{{NS['w']}}}left"),
                        'hanging': ind.get(f"{{{NS['w']}}}hanging"),
                    }
            levels.append(entry)
        abstract_defs[abstract_id] = {'levels': levels}

    numbering: dict[str, Any] = {}
    for num in root.findall('.//w:num', NS):
        num_id = num.get(f"{{{NS['w']}}}numId")
        abstract_num_id = num.find('w:abstractNumId', NS)
        if num_id and abstract_num_id is not None:
            abstract_id = abstract_num_id.get(f"{{{NS['w']}}}val")
            numbering[num_id] = {
                'abstractNumId': abstract_id,
                'levels': deepcopy(abstract_defs.get(abstract_id, {}).get('levels', [])),
            }
    return numbering


def _tbl_borders(tbl_pr: ET.Element | None) -> dict[str, Any]:
    if tbl_pr is None:
        return {}
    borders = tbl_pr.find('w:tblBorders', NS)
    if borders is None:
        return {}
    result = {}
    for edge in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
        el = borders.find(f'w:{edge}', NS)
        if el is not None:
            result[edge] = {
                'val': el.get(f"{{{NS['w']}}}val"),
                'sz': el.get(f"{{{NS['w']}}}sz"),
                'color': el.get(f"{{{NS['w']}}}color"),
            }
    return result


def _parse_tables_xml(root: ET.Element) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    for idx, tbl in enumerate(root.findall('.//w:tbl', NS)):
        tbl_pr = tbl.find('w:tblPr', NS)
        grid = tbl.find('w:tblGrid', NS)
        rows = tbl.findall('w:tr', NS)
        info: dict[str, Any] = {
            'index': idx,
            'rows': len(rows),
            'properties': {},
            'grid': [],
            'sample_cells': [],
        }
        if tbl_pr is not None:
            tbl_w = tbl_pr.find('w:tblW', NS)
            if tbl_w is not None:
                info['properties']['width'] = {
                    'value': tbl_w.get(f"{{{NS['w']}}}w"),
                    'type': tbl_w.get(f"{{{NS['w']}}}type"),
                }
            tbl_style = tbl_pr.find('w:tblStyle', NS)
            if tbl_style is not None:
                info['properties']['style'] = tbl_style.get(f"{{{NS['w']}}}val")
            tbl_layout = tbl_pr.find('w:tblLayout', NS)
            if tbl_layout is not None:
                info['properties']['layout'] = tbl_layout.get(f"{{{NS['w']}}}type")
            info['properties']['borders'] = _tbl_borders(tbl_pr)
        if grid is not None:
            info['grid'] = [col.get(f"{{{NS['w']}}}w") for col in grid.findall('w:gridCol', NS)]
        for r_idx, tr in enumerate(rows[:4]):
            cells = []
            for c_idx, tc in enumerate(tr.findall('w:tc', NS)[:6]):
                texts = [normalize_space(''.join(t.itertext())) for t in tc.findall('.//w:t', NS)]
                tcPr = tc.find('w:tcPr', NS)
                cell_info = {'row': r_idx, 'col': c_idx, 'text': normalize_space(' '.join(t for t in texts if t))}
                if tcPr is not None:
                    tcw = tcPr.find('w:tcW', NS)
                    if tcw is not None:
                        cell_info['width'] = tcw.get(f"{{{NS['w']}}}w")
                    grid_span = tcPr.find('w:gridSpan', NS)
                    if grid_span is not None:
                        cell_info['gridSpan'] = grid_span.get(f"{{{NS['w']}}}val")
                    vmerge = tcPr.find('w:vMerge', NS)
                    if vmerge is not None:
                        cell_info['vMerge'] = vmerge.get(f"{{{NS['w']}}}val") or 'continue'
                cells.append(cell_info)
            info['sample_cells'].append(cells)
        tables.append(info)
    return tables


def _parse_sections_xml(document_root: ET.Element) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    sect_paths = document_root.findall('.//w:sectPr', NS)
    for idx, sect in enumerate(sect_paths):
        pgSz = sect.find('w:pgSz', NS)
        pgMar = sect.find('w:pgMar', NS)
        cols = sect.find('w:cols', NS)
        titlePg = sect.find('w:titlePg', NS)
        refs = []
        for ref in sect.findall('w:headerReference', NS) + sect.findall('w:footerReference', NS):
            refs.append({
                'tag': ref.tag.split('}')[-1],
                'type': ref.get(f"{{{NS['w']}}}type"),
                'id': ref.get(f"{{{NS['r']}}}id"),
            })
        sections.append({
            'index': idx,
            'page_size': {
                'w': pgSz.get(f"{{{NS['w']}}}w") if pgSz is not None else None,
                'h': pgSz.get(f"{{{NS['w']}}}h") if pgSz is not None else None,
                'orient': pgSz.get(f"{{{NS['w']}}}orient") if pgSz is not None else None,
            },
            'page_margins': {
                'top': pgMar.get(f"{{{NS['w']}}}top") if pgMar is not None else None,
                'right': pgMar.get(f"{{{NS['w']}}}right") if pgMar is not None else None,
                'bottom': pgMar.get(f"{{{NS['w']}}}bottom") if pgMar is not None else None,
                'left': pgMar.get(f"{{{NS['w']}}}left") if pgMar is not None else None,
                'header': pgMar.get(f"{{{NS['w']}}}header") if pgMar is not None else None,
                'footer': pgMar.get(f"{{{NS['w']}}}footer") if pgMar is not None else None,
                'gutter': pgMar.get(f"{{{NS['w']}}}gutter") if pgMar is not None else None,
            },
            'cols': cols.get(f"{{{NS['w']}}}num") if cols is not None else None,
            'titlePg': titlePg is not None,
            'references': refs,
        })
    return sections


def _parse_theme_xml(root: ET.Element) -> dict[str, Any]:
    theme = {'latin_typefaces': {}, 'east_asia_typefaces': {}}
    major = root.find('.//a:themeElements/a:fontScheme/a:majorFont', NS)
    minor = root.find('.//a:themeElements/a:fontScheme/a:minorFont', NS)
    if major is not None:
        latin = major.find('a:latin', NS)
        ea = major.find('a:ea', NS)
        theme['latin_typefaces']['major'] = latin.get('typeface') if latin is not None else None
        theme['east_asia_typefaces']['major'] = ea.get('typeface') if ea is not None else None
    if minor is not None:
        latin = minor.find('a:latin', NS)
        ea = minor.find('a:ea', NS)
        theme['latin_typefaces']['minor'] = latin.get('typeface') if latin is not None else None
        theme['east_asia_typefaces']['minor'] = ea.get('typeface') if ea is not None else None
    return theme


def _ooxml_summary(docx_path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {
        'styles': {},
        'numbering': {},
        'tables': [],
        'headers_footers': {'headers': [], 'footers': []},
        'sections': [],
        'theme': {},
        'text_parts': [],
    }
    with zipfile.ZipFile(docx_path, 'r') as zf:
        names = set(zf.namelist())
        if 'word/styles.xml' in names:
            with zf.open('word/styles.xml') as f:
                result['styles'] = _parse_styles_xml(ET.parse(f).getroot())
        if 'word/numbering.xml' in names:
            with zf.open('word/numbering.xml') as f:
                result['numbering'] = _parse_numbering_xml(ET.parse(f).getroot())
        if 'word/document.xml' in names:
            with zf.open('word/document.xml') as f:
                root = ET.parse(f).getroot()
                result['tables'] = _parse_tables_xml(root)
                result['sections'] = _parse_sections_xml(root)
        if 'word/theme/theme1.xml' in names:
            with zf.open('word/theme/theme1.xml') as f:
                result['theme'] = _parse_theme_xml(ET.parse(f).getroot())
        for name in sorted(names):
            if name.startswith('word/header') and name.endswith('.xml'):
                result['headers_footers']['headers'].append({'file': name, 'exists': True})
            elif name.startswith('word/footer') and name.endswith('.xml'):
                result['headers_footers']['footers'].append({'file': name, 'exists': True})
            if name.startswith('word/') and name.endswith('.xml') and ('header' in name or 'footer' in name or name in OOXML_TEXT_PARTS):
                result['text_parts'].append(name)
    return result


def structural_signature(profile: dict[str, Any], limit: int = 16) -> list[dict[str, Any]]:
    sig: list[dict[str, Any]] = []
    for e in profile.get('paragraphs_preview', [])[:limit]:
        sig.append({
            'centered': e.get('alignment') == WD_ALIGN_PARAGRAPH.CENTER,
            'all_caps': e.get('all_caps', False),
            'has_so': e.get('text', '').startswith('Số:'),
            'has_kinh_gui': 'Kính gửi' in e.get('text', ''),
            'has_noi_nhan': 'Nơi nhận' in e.get('text', ''),
            'origin': e.get('origin', ''),
            'style_name': e.get('style_name', ''),
            'jc': e.get('xml', {}).get('jc'),
            'has_tabs': bool(e.get('xml', {}).get('tabs')),
            'has_indent': bool(e.get('xml', {}).get('indent')),
        })
    return sig


def extract_template_profile(path: str | Path) -> dict[str, Any]:
    path = Path(path)
    doc = Document(path)
    entries = paragraph_entries(doc)
    all_text = '\n'.join(e['text'] for e in entries)
    style_names: list[str] = []
    for st in doc.styles:
        try:
            name = getattr(st, 'name', '')
            if name:
                style_names.append(name)
        except Exception:
            pass
    placeholders: list[str] = []
    for e in entries:
        placeholders.extend(PLACEHOLDER_RE.findall(e['text']))
    text_meta = header_footer_text(doc)
    ooxml = _ooxml_summary(path)
    flags = {
        'has_quoc_hieu': 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM' in all_text.upper(),
        'has_tieu_ngu': 'ĐỘC LẬP - TỰ DO - HẠNH PHÚC' in all_text.upper(),
        'has_number_line': bool(NUMBER_LINE_RE.search(all_text)),
        'has_date_line': bool(DATE_LINE_RE.search(all_text)),
        'has_noi_nhan_or_luu': ('NƠI NHẬN' in all_text.upper()) or ('LƯU:' in all_text.upper()),
    }
    titles = title_candidates(entries)
    profile: dict[str, Any] = {
        'profile_version': '2.1',
        'source_document': path.name,
        'source_path': str(path),
        'detected_document_type': detect_document_type(all_text, entries),
        'section_metrics': section_metrics(doc),
        'headers': text_meta['headers'],
        'footers': text_meta['footers'],
        'table_count': len(doc.tables),
        'style_names': sorted(set(style_names)),
        'paragraphs_preview': entries[:80],
        'nonempty_paragraph_count': len(entries),
        'placeholder_keys': sorted(set(placeholders)),
        'title_candidates': titles,
        'named_title_candidate': guess_named_document_title(titles),
        'key_lines': [e['text'] for e in entries[:30]],
        'core_flags': flags,
        'has_quoc_hieu': flags['has_quoc_hieu'],
        'has_tieu_ngu': flags['has_tieu_ngu'],
        'has_number_line': flags['has_number_line'],
        'has_date_line': flags['has_date_line'],
        'styles': ooxml['styles'],
        'numbering': ooxml['numbering'],
        'tables': ooxml['tables'],
        'headers_footers': ooxml['headers_footers'],
        'sections': ooxml['sections'],
        'theme': ooxml['theme'],
        'text_parts': ooxml['text_parts'],
    }
    profile['structural_signature'] = structural_signature(profile)
    return profile


def compare_profiles(candidate: dict[str, Any], template: dict[str, Any]) -> dict[str, Any]:
    diffs: list[str] = []
    score = 100.0
    cs = candidate.get('section_metrics', {})
    ts = template.get('section_metrics', {})
    for k in ['top_margin_mm', 'bottom_margin_mm', 'left_margin_mm', 'right_margin_mm']:
        cv, tv = cs.get(k), ts.get(k)
        if cv is not None and tv is not None and abs(cv - tv) > 3:
            diffs.append(f'{k} differs noticeably: candidate={cv} mm, template={tv} mm')
            score -= 8
    ctype = candidate.get('detected_document_type')
    ttype = template.get('detected_document_type')
    if ctype != ttype:
        diffs.append(f'detected type differs: candidate={ctype}, template={ttype}')
        score -= 12
    csig = candidate.get('structural_signature', [])
    tsig = template.get('structural_signature', [])
    if len(csig) != len(tsig):
        diffs.append(f'top-structure length differs: candidate={len(csig)} blocks, template={len(tsig)} blocks')
        score -= 8
    elif csig:
        mismatches = sum(
            1 for a, b in zip(csig, tsig)
            if (a.get('centered'), a.get('all_caps'), a.get('has_so'), a.get('has_kinh_gui'), a.get('origin'), a.get('jc'))
            != (b.get('centered'), b.get('all_caps'), b.get('has_so'), b.get('has_kinh_gui'), b.get('origin'), b.get('jc'))
        )
        if mismatches > 0:
            diffs.append(f'top structural signature differs in {mismatches}/{len(csig)} compared blocks')
            score -= min(18, mismatches * 1.5)
    if abs(candidate.get('table_count', 0) - template.get('table_count', 0)) > 1:
        diffs.append(f"table count differs: candidate={candidate.get('table_count')}, template={template.get('table_count')}")
        score -= 8
    cand_placeholders = set(candidate.get('placeholder_keys', []))
    tmpl_placeholders = set(template.get('placeholder_keys', []))
    if tmpl_placeholders and cand_placeholders != tmpl_placeholders:
        diffs.append('placeholder set differs from template')
        score -= 6
    cand_theme = candidate.get('theme', {})
    tmpl_theme = template.get('theme', {})
    if cand_theme and tmpl_theme and cand_theme != tmpl_theme:
        diffs.append('theme font scheme differs from template')
        score -= 4
    score = max(0.0, round(score, 1))
    return {'matched': not diffs, 'similarity_score': score, 'differences': diffs}


def flatten_for_placeholders(value: Any) -> str:
    if value is None:
        return ''
    if isinstance(value, str):
        return value
    if isinstance(value, (int, float, bool)):
        return str(value)
    if isinstance(value, list):
        if all(isinstance(x, dict) for x in value):
            lines = []
            for idx, item in enumerate(value, start=1):
                title = normalize_space(_safe(item.get('title')))
                body = item.get('body')
                body_text = '\n'.join(normalize_space(_safe(x)) for x in body) if isinstance(body, list) else normalize_space(_safe(body))
                label = f'Điều {idx}. {title}'.strip()
                lines.append('\n'.join(x for x in [label, body_text] if x))
            return '\n\n'.join(lines)
        return '\n'.join(normalize_space(_safe(x)) for x in value)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return _safe(value)


def content_to_placeholder_map(content: dict[str, Any]) -> dict[str, str]:
    mapping = {str(k): flatten_for_placeholders(v) for k, v in content.items() if k not in {'placeholder_values'}}
    custom = content.get('placeholder_values', {})
    if isinstance(custom, dict):
        for k, v in custom.items():
            mapping[str(k)] = flatten_for_placeholders(v)
    if 'issue_date' in mapping and re.fullmatch(r'\d{4}-\d{2}-\d{2}', mapping['issue_date']):
        y, m, d = mapping['issue_date'].split('-')
        mapping.setdefault('issue_date_text', f'ngày {d} tháng {m} năm {y}')
    if 'document_number' in content or 'document_symbol' in content:
        number = normalize_space(_safe(content.get('document_number')))
        symbol = normalize_space(_safe(content.get('document_symbol')))
        mapping.setdefault('so_ky_hieu', f'Số: {number}/{symbol}' if symbol else f'Số: {number}')
    if 'location' in mapping and 'issue_date_text' in mapping:
        mapping.setdefault('date_line', f"{mapping['location']}, {mapping['issue_date_text']}")
    if 'parent_agency' in mapping and 'issuing_agency' in mapping:
        mapping.setdefault('co_quan_2_dong', '\n'.join(x for x in [mapping['parent_agency'], mapping['issuing_agency']] if x))
    if 'kinh_gui' in content and isinstance(content['kinh_gui'], list):
        recips = [normalize_space(_safe(x)) for x in content['kinh_gui'] if normalize_space(_safe(x))]
        mapping.setdefault('kinh_gui_block', '\n'.join(f'- {x};' if i < len(recips) - 1 else f'- {x}.' for i, x in enumerate(recips)))
    if 'noi_nhan' in content and isinstance(content['noi_nhan'], list):
        recips = [normalize_space(_safe(x)) for x in content['noi_nhan'] if normalize_space(_safe(x))]
        mapping.setdefault('noi_nhan_block', '\n'.join(f'- {x}' for x in recips))
    if 'body_paragraphs' in content and isinstance(content['body_paragraphs'], list):
        mapping.setdefault('body_block', '\n\n'.join(normalize_space(_safe(x)) for x in content['body_paragraphs']))
    return mapping


def _replace_in_paragraph_xml(paragraph: ET.Element, data: dict[str, str]) -> int:
    text_nodes = paragraph.findall('.//w:t', NS)
    if not text_nodes:
        return 0
    full = ''.join(node.text or '' for node in text_nodes)
    if '{{' not in full or '}}' not in full:
        return 0
    replaced = full
    lower = {k.lower(): v for k, v in data.items()}

    def repl(match: re.Match[str]) -> str:
        key = match.group(1)
        return lower.get(key.lower(), data.get(key, ''))

    replaced = PLACEHOLDER_RE.sub(repl, replaced)
    if replaced == full:
        return 0
    text_nodes[0].text = replaced
    for node in text_nodes[1:]:
        node.text = ''
    return 1


def clone_patch_docx(source_docx: Path, output_docx: Path, data: dict[str, str]) -> dict[str, Any]:
    patched_parts: dict[str, int] = {}
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_docx, output_docx)
    with TemporaryDirectory(prefix='nd30_clone_patch_') as tmpdir:
        tmp = Path(tmpdir)
        with zipfile.ZipFile(output_docx, 'r') as zf:
            zf.extractall(tmp)
        names = [p.relative_to(tmp).as_posix() for p in tmp.rglob('*.xml')]
        for rel in names:
            if not (rel.startswith('word/header') or rel.startswith('word/footer') or rel in OOXML_TEXT_PARTS):
                continue
            xml_path = tmp / rel
            try:
                root = ET.parse(xml_path).getroot()
            except ET.ParseError:
                continue
            count = 0
            for para in root.findall('.//w:p', NS):
                count += _replace_in_paragraph_xml(para, data)
            if count:
                ET.ElementTree(root).write(xml_path, encoding='utf-8', xml_declaration=True)
                patched_parts[rel] = count
        with zipfile.ZipFile(output_docx, 'w', zipfile.ZIP_DEFLATED) as zf:
            for file_path in tmp.rglob('*'):
                if file_path.is_file():
                    zf.write(file_path, file_path.relative_to(tmp))
    return {
        'patched_parts': patched_parts,
        'replacement_count': sum(patched_parts.values()),
        'mode': 'clone-patch',
    }


def dump_json(data: dict[str, Any], path: str | Path) -> None:
    Path(path).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')
