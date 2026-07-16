#!/usr/bin/env python3
"""Create ND30-oriented Vietnamese administrative DOCX files from JSON input.

Usage:
    python create_nd30_docx.py input.json output.docx
"""
from __future__ import annotations

import json
import re
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


@dataclass
class Spec:
    data: dict[str, Any]

    def get(self, key: str, default: Any = "") -> Any:
        return self.data.get(key, default)

    def list(self, key: str) -> list[Any]:
        value = self.data.get(key, [])
        return value if isinstance(value, list) else []


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    ensure_page_numbers(doc)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = f"w:{edge}"
            el = tcBorders.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                tcBorders.append(el)
            for key in ("val", "sz", "space", "color"):
                if key in edge_data:
                    el.set(qn(f"w:{key}"), str(edge_data[key]))


def add_run(paragraph, text: str, *, bold=False, italic=False, size=14, uppercase=False):
    run = paragraph.add_run((text or "").upper() if uppercase else (text or ""))
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    return run

def set_paragraph_bottom_border(paragraph, size: int = 8, color: str = "000000"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_rule_paragraph(container, *, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(0), space_after=Pt(2)):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_bottom_border(p)
    return p


def three_zone_table(doc: Document, widths=(7.2, 1.2, 8.0), rows: int = 1):
    table = doc.add_table(rows=rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, w in enumerate(widths):
        table.columns[i].width = Cm(w)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell, top={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"}, bottom={"val": "nil"})
    return table



def _append_page_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    for node in (begin, instr, sep, txt, end):
        run._r.append(node)


def ensure_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = True
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        if 'PAGE' not in p._p.xml:
            _append_page_field(p)
        first = section.first_page_header
        if first.paragraphs:
            first.paragraphs[0].text = ''


def set_run_tracking(paragraph, value: int = -2) -> None:
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        spacing = rPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            rPr.append(spacing)
        spacing.set(qn('w:val'), str(value))

def _set_on_off(pPr, tag: str, enabled: bool = True):
    if pPr is None or not enabled:
        return
    el = pPr.find(qn(f"w:{tag}"))
    if el is None:
        el = OxmlElement(f"w:{tag}")
        pPr.append(el)


def apply_pagination_controls(paragraph, *, role: str = 'body', keep_next: bool = False, keep_lines: bool = False, widow_control: bool = True):
    pPr = paragraph._p.get_or_add_pPr()
    _set_on_off(pPr, 'widowControl', widow_control)
    _set_on_off(pPr, 'keepNext', keep_next)
    _set_on_off(pPr, 'keepLines', keep_lines)
    if role in {'title', 'heading', 'heading-major', 'heading-minor', 'label'}:
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.0
    elif role == 'body':
        paragraph.paragraph_format.line_spacing = 1.1
        paragraph.paragraph_format.space_after = Pt(6)


def classify_text_role(text: str) -> str:
    raw = (text or '').strip()
    upper = raw.upper()
    if not raw:
        return 'body'
    if upper in {'NGHỊ QUYẾT', 'QUYẾT ĐỊNH', 'CÔNG ĐIỆN', 'GIẤY MỜI', 'GIẤY GIỚI THIỆU', 'BIÊN BẢN', 'GIẤY NGHỈ PHÉP'}:
        return 'title'
    if re.match(r'^(PHẦN|CHƯƠNG)\s+[IVXLC]+$', upper):
        return 'heading-major'
    if re.match(r'^(MỤC|TIỂU MỤC)\s+\d+', upper):
        return 'heading-major'
    if re.match(r'^[IVXLC]+\.\s+', upper):
        return 'heading-major'
    if re.match(r'^ĐIỀU\s+\d+\.', upper):
        return 'heading-minor'
    if re.match(r'^\d+\.\s+', raw):
        return 'heading-minor'
    if re.match(r'^[a-zà-ỹ]\)\s+', raw, re.I):
        return 'label'
    if raw.endswith(':') and len(raw) < 80:
        return 'label'
    return 'body'


def fmt_date(date_text: str) -> str:
    if not date_text:
        return "ngày ... tháng ... năm ..."
    try:
        dt = datetime.fromisoformat(date_text)
        return f"ngày {dt.day:02d} tháng {dt.month:02d} năm {dt.year}"
    except ValueError:
        return date_text


def para(doc: Document, text: str = "", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         first_line_cm: float | None = 1.0, bold=False, italic=False, size=14,
         space_after=Pt(6), space_before=Pt(0), role: str | None = None,
         keep_next: bool = False, keep_lines: bool = False, widow_control: bool = True):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Cm(first_line_cm) if first_line_cm is not None else None
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = 1.1
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size)
    apply_pagination_controls(p, role=role or classify_text_role(text), keep_next=keep_next, keep_lines=keep_lines, widow_control=widow_control)
    return p


def left_right_table(doc: Document, left_w=8.5, right_w=8.0):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(left_w)
    table.columns[1].width = Cm(right_w)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_border(cell, top={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"}, bottom={"val": "nil"})
    return left, right



def header_block(doc: Document, spec: Spec):
    table = three_zone_table(doc, rows=1)
    left = table.cell(0, 0)
    right = table.cell(0, 2)

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    parent = spec.get('parent_agency', '')
    if parent:
        add_run(p, parent, size=12, uppercase=True)
    p2 = left.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    add_run(p2, spec.get('issuing_agency', ''), bold=True, size=12, uppercase=True)
    add_rule_paragraph(left)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', bold=True, size=12)
    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    add_run(p2, 'Độc lập - Tự do - Hạnh phúc', bold=True, size=13)
    add_rule_paragraph(right)


def number_date_line(doc: Document, spec: Spec, *, inline_subject=False):
    left, right = left_right_table(doc)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    number = str(spec.get("document_number", "")).strip()
    symbol = str(spec.get("document_symbol", "")).strip()
    line = f"Số: {number}/{symbol}" if symbol else f"Số: {number}"
    add_run(p, line.strip(), size=13)
    if inline_subject and spec.get("subject"):
        p2 = left.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(6)
        p2.paragraph_format.space_after = Pt(0)
        add_run(p2, spec.get("subject"), size=12)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f"{spec.get('location', '...')}, {fmt_date(spec.get('issue_date', ''))}", italic=True, size=13)



def title_block(doc: Document, title: str, subject: str = ''):
    p = para(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_before=Pt(12), space_after=Pt(0), role='title', keep_next=True, keep_lines=True)
    add_run(p, title, bold=True, size=14, uppercase=True)
    if subject:
        p2 = para(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_after=Pt(0), role='heading', keep_next=True, keep_lines=True)
        add_run(p2, subject, bold=True, size=14)
    t = three_zone_table(doc, widths=(5.4, 5.4, 5.4), rows=1)
    center = t.cell(0, 1)
    add_rule_paragraph(center)



def add_kinh_gui(doc: Document, recipients: Iterable[str]):
    recipients = [str(x).strip() for x in recipients if str(x).strip()]
    if not recipients:
        return
    table = three_zone_table(doc, rows=1)
    label = table.cell(0, 0)
    body = table.cell(0, 1).merge(table.cell(0, 2))
    p = label.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    add_run(p, 'Kính gửi:', size=14)
    first = True
    for item in recipients:
        q = body.paragraphs[0] if first and not body.paragraphs[0].text else body.add_paragraph()
        first = False
        q.alignment = WD_ALIGN_PARAGRAPH.LEFT
        q.paragraph_format.space_after = Pt(0)
        q.paragraph_format.line_spacing = 1.0
        suffix = '.' if item == recipients[-1] else ';'
        add_run(q, f'- {item}{suffix}', size=14)



def add_body_paragraphs(doc: Document, paragraphs: list[Any]):
    for item in paragraphs:
        if isinstance(item, dict):
            item_type = str(item.get('type', 'paragraph')).strip().lower()
            if item_type == 'page_break':
                doc.add_page_break()
            else:
                para(doc, str(item.get('text', '')), role=str(item.get('role') or 'body'), keep_next=bool(item.get('keep_next', False)))
        else:
            para(doc, str(item), role='body')


def add_legal_bases(doc: Document, bases: list[str], authority_line: str = ""):
    if authority_line:
        p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_after=Pt(6))
        add_run(p, authority_line, bold=True, size=14, uppercase=True)
    for i, line in enumerate([str(x).strip() for x in bases if str(x).strip()]):
        if not line.endswith((";", ".")):
            line = line + ("." if i == len(bases) - 1 else ";")
        para(doc, line, italic=True)


def add_articles(doc: Document, articles: list[dict[str, Any]]):
    for idx, article in enumerate(articles, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(6)
        add_run(p, f"Điều {idx}. ", bold=True, size=14)
        add_run(p, str(article.get("title", "")), bold=True, size=14)
        apply_pagination_controls(p, role='heading-minor', keep_next=True, keep_lines=True)
        body = article.get("body", [])
        for line in body if isinstance(body, list) else [str(body)]:
            para(doc, str(line), role='body')



def signature_block(doc: Document, spec: Spec):
    table = three_zone_table(doc, widths=(6.8, 1.4, 7.5), rows=1)
    left = table.cell(0, 0)
    right = table.cell(0, 2)

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, 'Nơi nhận:', bold=True, italic=True, size=12)
    items = [str(x).strip() for x in spec.list('noi_nhan') if str(x).strip()]
    if not items:
        items = ['Lưu: VT.']
    for item in items:
        q = left.add_paragraph()
        q.paragraph_format.space_after = Pt(0)
        q.paragraph_format.line_spacing = 1.0
        add_run(q, f'- {item}', size=11)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    auth = str(spec.get('signer_authority', '')).strip()
    behalf = str(spec.get('signer_on_behalf_of', '')).strip()
    position = str(spec.get('signer_position', '')).strip()
    second_position = str(spec.get('signer_second_position', '')).strip()
    signer_name = str(spec.get('signer_name', '')).strip()
    if auth and behalf:
        add_run(p, f'{auth} {behalf}', bold=True, size=14, uppercase=True)
    elif auth and position:
        add_run(p, f'{auth} {position}', bold=True, size=14, uppercase=True)
        position = ''
    elif auth:
        add_run(p, auth, bold=True, size=14, uppercase=True)
    if position:
        p2 = right.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.0
        add_run(p2, position, bold=True, size=14, uppercase=True)
    if second_position:
        p3 = right.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        p3.paragraph_format.line_spacing = 1.0
        add_run(p3, second_position, bold=True, size=14, uppercase=True)
    gap = right.add_paragraph()
    gap.paragraph_format.space_before = Pt(34)
    gap.paragraph_format.space_after = Pt(0)
    gap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p = right.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    add_run(name_p, signer_name, bold=True, size=14)


def build_cong_van(doc: Document, spec: Spec):
    header_block(doc, spec)
    number_date_line(doc, spec, inline_subject=True)
    add_kinh_gui(doc, spec.list("kinh_gui"))
    add_body_paragraphs(doc, spec.list("body_paragraphs"))
    signature_block(doc, spec)


def build_nghi_quyet(doc: Document, spec: Spec):
    header_block(doc, spec)
    number_date_line(doc, spec)
    title_block(doc, "NGHỊ QUYẾT", spec.get("subject", ""))
    add_legal_bases(doc, spec.list("legal_bases"), spec.get("authority_line", spec.get("issuing_agency", "")))
    p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_after=Pt(6), role='heading', keep_next=True, keep_lines=True)
    add_run(p, "QUYẾT NGHỊ:", bold=True, size=14, uppercase=True)
    add_articles(doc, spec.list("articles")) if spec.list("articles") else add_body_paragraphs(doc, spec.list("body_paragraphs"))
    signature_block(doc, spec)


def build_named_doc(doc: Document, spec: Spec):
    header_block(doc, spec)
    number_date_line(doc, spec)
    title_block(doc, spec.get("document_title", ""), spec.get("subject", ""))
    intro = str(spec.get("intro", "")).strip()
    if intro:
        para(doc, intro)
    add_body_paragraphs(doc, spec.list("body_paragraphs"))
    signature_block(doc, spec)


def build_quyet_dinh(doc: Document, spec: Spec):
    header_block(doc, spec)
    number_date_line(doc, spec)
    title_block(doc, "QUYẾT ĐỊNH", spec.get("subject", ""))
    add_legal_bases(doc, spec.list("legal_bases"), spec.get("authority_line", spec.get("signer_position", "")))
    p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_after=Pt(6), role='heading', keep_next=True, keep_lines=True)
    add_run(p, "QUYẾT ĐỊNH:", bold=True, size=14)
    if spec.list("articles"):
        add_articles(doc, spec.list("articles"))
    else:
        add_body_paragraphs(doc, spec.list("body_paragraphs"))
    signature_block(doc, spec)


def build_cong_dien(doc: Document, spec: Spec):
    header_block(doc, spec)
    number_date_line(doc, spec)
    title_block(doc, "CÔNG ĐIỆN", spec.get("subject", ""))
    if spec.get("sender_authority"):
        p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_after=Pt(6))
        add_run(p, f"{spec.get('sender_authority')} điện:", size=14)
    add_kinh_gui(doc, spec.list("kinh_gui"))
    add_body_paragraphs(doc, spec.list("body_paragraphs"))
    signature_block(doc, spec)


def build_giay_moi(doc: Document, spec: Spec):
    header_block(doc, spec)
    number_date_line(doc, spec)
    title_block(doc, "GIẤY MỜI", spec.get("subject", ""))
    para(doc, f"{spec.get('issuing_agency')} trân trọng kính mời: {spec.get('invitee')}", first_line_cm=None, role='label', keep_next=True, keep_lines=True)
    para(doc, f"Tới dự: {spec.get('event_name')}", first_line_cm=None)
    para(doc, f"Chủ trì: {spec.get('event_chair')}", first_line_cm=None)
    para(doc, f"Thời gian: {spec.get('event_time')}", first_line_cm=None)
    para(doc, f"Địa điểm: {spec.get('event_location')}", first_line_cm=None)
    if spec.get('event_notes'):
        para(doc, str(spec.get('event_notes')), first_line_cm=None)
    signature_block(doc, spec)


def build_giay_gioi_thieu(doc: Document, spec: Spec):
    header_block(doc, spec)
    number_date_line(doc, spec)
    title_block(doc, "GIẤY GIỚI THIỆU")
    para(doc, f"{spec.get('issuing_agency')} trân trọng giới thiệu:", first_line_cm=None)
    para(doc, f"Ông (bà): {spec.get('introduced_person')}", first_line_cm=None)
    if spec.get('introduced_role'):
        para(doc, f"Chức vụ: {spec.get('introduced_role')}", first_line_cm=None)
    para(doc, f"Được cử đến: {spec.get('destination_agency')}", first_line_cm=None)
    para(doc, f"Về việc: {spec.get('purpose')}", first_line_cm=None)
    if spec.get('valid_until'):
        para(doc, f"Giấy này có giá trị đến hết ngày {spec.get('valid_until')}./.", first_line_cm=None)
    else:
        para(doc, "Đề nghị Quý cơ quan tạo điều kiện để người có tên ở trên hoàn thành nhiệm vụ./.", first_line_cm=None)
    signature_block(doc, spec)


def build_bien_ban(doc: Document, spec: Spec):
    header_block(doc, spec)
    title_block(doc, "BIÊN BẢN", spec.get("subject", ""))
    for label, key in [
        ('Thời gian bắt đầu', 'meeting_start'),
        ('Địa điểm', 'meeting_location'),
        ('Thành phần tham dự', 'meeting_attendees'),
        ('Chủ trì (chủ tọa)', 'meeting_chair'),
        ('Thư ký (người ghi biên bản)', 'meeting_secretary'),
    ]:
        para(doc, f"{label}: {spec.get(key)}", first_line_cm=None)
    para(doc, 'Nội dung (theo diễn biến cuộc họp/hội nghị/hội thảo):', first_line_cm=None, role='label', keep_next=True, keep_lines=True)
    add_body_paragraphs(doc, spec.list('body_paragraphs'))
    if spec.get('meeting_end'):
        para(doc, f"Cuộc họp kết thúc vào {spec.get('meeting_end')}./.", first_line_cm=None)
    left, right = left_right_table(doc)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'THƯ KÝ', bold=True, size=14, uppercase=True)
    p.add_run('\n\n\n\n')
    add_run(p, spec.get('secretary_name', spec.get('meeting_secretary', '')), bold=True, size=14)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, 'CHỦ TỌA', bold=True, size=14, uppercase=True)
    p.add_run('\n\n\n\n')
    add_run(p, spec.get('chair_name', spec.get('meeting_chair', '')), bold=True, size=14)
    if spec.list('noi_nhan'):
        para(doc, '', first_line_cm=None, space_after=Pt(0))
        signature_block(doc, Spec({**spec.data, 'signer_name': '', 'signer_position': ''}))


def build_giay_nghi_phep(doc: Document, spec: Spec):
    header_block(doc, spec)
    number_date_line(doc, spec)
    title_block(doc, 'GIẤY NGHỈ PHÉP')
    if spec.get('leave_request_date') and spec.get('leave_requester'):
        para(doc, f"Xét Đơn đề nghị nghỉ phép ngày {spec.get('leave_request_date')} của ông (bà) {spec.get('leave_requester')}", first_line_cm=None)
    para(doc, f"{spec.get('issuing_agency')} cấp cho:", first_line_cm=None)
    para(doc, f"Ông (bà): {spec.get('leave_person')}", first_line_cm=None)
    if spec.get('leave_role'):
        para(doc, f"Chức vụ: {spec.get('leave_role')}", first_line_cm=None)
    para(doc, f"Được nghỉ phép trong thời gian {spec.get('leave_period')} tại {spec.get('leave_location')}", first_line_cm=None)
    if spec.get('leave_basis'):
        para(doc, f"Số ngày nghỉ phép nêu trên được tính vào thời gian {spec.get('leave_basis')}./.", first_line_cm=None)
    signature_block(doc, spec)


def build_phu_luc(doc: Document, spec: Spec):
    p = para(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_after=Pt(0))
    add_run(p, f"Phụ lục {spec.get('appendix_number')}", bold=True, size=14)
    p2 = para(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_after=Pt(0))
    add_run(p2, spec.get('appendix_title', ''), bold=True, size=14, uppercase=True)
    p3 = para(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_after=Pt(6))
    add_run(p3, f"({spec.get('appendix_parent_note')})", italic=True, size=13)
    add_body_paragraphs(doc, spec.list('body_paragraphs'))


def build_ban_sao(doc: Document, spec: Spec, *, electronic: bool):
    p = para(doc, '', align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, space_after=Pt(0))
    add_run(p, spec.get('copy_form', ''), bold=True, size=14, uppercase=True)
    left, right = left_right_table(doc)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, spec.get('copy_agency', ''), bold=True, size=12, uppercase=True)
    p.add_run('\n')
    copy_symbol = str(spec.get('copy_symbol', '')).strip()
    number = str(spec.get('copy_number', '')).strip()
    add_run(p, f"Số: {number}/{copy_symbol}" if copy_symbol else f"Số: {number}", size=13)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"{spec.get('location', '...')}, {fmt_date(spec.get('issue_date', ''))}", italic=True, size=13)
    if electronic:
        para(doc, f"{spec.get('copy_agency')} ký số bản sao điện tử theo quy định./.", first_line_cm=None)
    signature_block(doc, spec)




def optimize_pagination(doc: Document) -> None:
    paragraphs = [p for p in doc.paragraphs if (p.text or '').strip()]
    for idx, p in enumerate(paragraphs):
        role = classify_text_role(p.text)
        keep_next = role in {'title', 'heading-major', 'heading-minor', 'label'}
        keep_lines = role in {'title', 'heading-major', 'heading-minor'} or (role == 'body' and len((p.text or '').split()) <= 10)
        apply_pagination_controls(p, role=role, keep_next=keep_next, keep_lines=keep_lines, widow_control=True)
        words = len((p.text or '').split())
        if role == 'body' and words > 90:
            p.paragraph_format.line_spacing = 1.02
            p.paragraph_format.space_after = Pt(4)
            set_run_tracking(p, -4)
        elif role == 'body' and words > 55:
            p.paragraph_format.line_spacing = 1.04
            p.paragraph_format.space_after = Pt(4)
            set_run_tracking(p, -3)
        elif role == 'body' and words > 35:
            p.paragraph_format.line_spacing = 1.06
            p.paragraph_format.space_after = Pt(5)
            set_run_tracking(p, -2)


def add_box_mark(doc: Document, text: str):
    text = (text or "").strip()
    if not text:
        return
    p = para(doc, "", first_line_cm=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(0))
    add_run(p, f"[{text}]", bold=True, size=13, uppercase=True)


def contact_line(doc: Document, spec: Spec):
    contact = str(spec.get("contact_line", "")).strip()
    if not contact:
        return
    p = para(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, size=11, space_before=Pt(6), space_after=Pt(0))
    set_paragraph_bottom_border(p, size=4)
    p2 = para(doc, contact, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=None, size=11, space_after=Pt(0))


def drafter_mark(doc: Document, spec: Spec):
    mark = str(spec.get("drafter_mark", "")).strip()
    if mark:
        para(doc, mark, first_line_cm=None, size=11, space_before=Pt(3), space_after=Pt(0))

def build(doc: Document, spec: Spec):
    add_box_mark(doc, spec.get('urgency', ''))
    add_box_mark(doc, spec.get('confidentiality', ''))
    add_box_mark(doc, spec.get('circulation_scope', ''))
    kind = spec.get('document_type', 'cong-van')
    if kind == 'cong-van':
        build_cong_van(doc, spec)
    elif kind == 'nghi-quyet-ca-biet':
        build_nghi_quyet(doc, spec)
    elif kind in {'quyet-dinh-truc-tiep', 'quyet-dinh-gian-tiep'}:
        build_quyet_dinh(doc, spec)
    elif kind == 'van-ban-co-ten-loai':
        build_named_doc(doc, spec)
    elif kind == 'cong-dien':
        build_cong_dien(doc, spec)
    elif kind == 'giay-moi':
        build_giay_moi(doc, spec)
    elif kind == 'giay-gioi-thieu':
        build_giay_gioi_thieu(doc, spec)
    elif kind == 'bien-ban':
        build_bien_ban(doc, spec)
    elif kind == 'giay-nghi-phep':
        build_giay_nghi_phep(doc, spec)
    elif kind == 'phu-luc':
        build_phu_luc(doc, spec)
    elif kind == 'ban-sao-giay':
        build_ban_sao(doc, spec, electronic=False)
    elif kind == 'ban-sao-dien-tu':
        build_ban_sao(doc, spec, electronic=True)
    else:
        raise SystemExit(f'Unsupported document_type: {kind}')
    contact_line(doc, spec)
    drafter_mark(doc, spec)


def load_spec(path: Path) -> Spec:
    try:
        data = json.loads(path.read_text(encoding='utf-8'))
        if not isinstance(data, dict):
            raise ValueError('Input JSON must be an object')
        return Spec(data)
    except Exception as exc:
        raise SystemExit(f'Failed to parse input JSON: {exc}') from exc


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print('Usage: python create_nd30_docx.py input.json output.docx')
        return 1
    input_path = Path(argv[1])
    output_path = Path(argv[2])
    if not input_path.exists():
        print(f'Input file not found: {input_path}')
        return 1
    spec = load_spec(input_path)
    doc = Document()
    set_doc_defaults(doc)
    build(doc, spec)
    optimize_pagination(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f'Created {output_path}')
    return 0


if __name__ == '__main__':
    raise SystemExit(main(sys.argv))
